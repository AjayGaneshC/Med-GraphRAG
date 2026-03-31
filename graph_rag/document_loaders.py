"""Document loaders for various file formats."""
from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load and extract text from various document formats."""
    
    @staticmethod
    def load_text(file_path: str) -> str:
        """Load plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def load_pdf(file_path: str) -> str:
        """Load PDF file and extract text."""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise
    
    @staticmethod
    def load_docx(file_path: str) -> str:
        """Load Word document (.docx) and extract text."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    
    @staticmethod
    def load_html(file_path: str) -> str:
        """Load HTML file and extract text."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'lxml')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text
            text = soup.get_text(separator="\n", strip=True)
            
            # Clean up whitespace
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)
        except ImportError:
            logger.error("beautifulsoup4 and lxml not installed. Install with: pip install beautifulsoup4 lxml")
            raise
        except Exception as e:
            logger.error(f"Error reading HTML: {e}")
            raise
    
    @staticmethod
    def load_csv(file_path: str) -> str:
        """Load CSV file and convert to text."""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text format
            text_parts = []
            text_parts.append(f"CSV Data with {len(df)} rows and {len(df.columns)} columns")
            text_parts.append(f"Columns: {', '.join(df.columns)}")
            text_parts.append("\n")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("Numeric Summary:")
                text_parts.append(df[numeric_cols].describe().to_string())
                text_parts.append("\n")
            
            # Add sample rows
            text_parts.append("Sample Data (first 100 rows):")
            text_parts.append(df.head(100).to_string(index=False))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            raise
    
    @staticmethod
    def load_json(file_path: str) -> str:
        """Load JSON file and convert to text."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Pretty print JSON with formatting
            json_str = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Also create a flattened text representation
            def extract_text_from_json(obj, prefix=""):
                """Recursively extract text from JSON."""
                texts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        new_prefix = f"{prefix}.{key}" if prefix else key
                        if isinstance(value, (dict, list)):
                            texts.extend(extract_text_from_json(value, new_prefix))
                        else:
                            texts.append(f"{new_prefix}: {value}")
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        new_prefix = f"{prefix}[{i}]"
                        if isinstance(item, (dict, list)):
                            texts.extend(extract_text_from_json(item, new_prefix))
                        else:
                            texts.append(f"{new_prefix}: {item}")
                return texts
            
            flat_text = "\n".join(extract_text_from_json(data))
            
            return f"JSON Structure:\n{json_str}\n\nFlattened Content:\n{flat_text}"
        except Exception as e:
            logger.error(f"Error reading JSON: {e}")
            raise
    
    @staticmethod
    def load_xml(file_path: str) -> str:
        """Load XML file and extract text."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f.read(), 'lxml-xml')
            
            # Get all text content
            text = soup.get_text(separator="\n", strip=True)
            
            # Clean up whitespace
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"Error reading XML: {e}")
            raise
    
    @staticmethod
    def load_rtf(file_path: str) -> str:
        """Load RTF file (basic extraction)."""
        # RTF is complex, but we can do basic text extraction
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Very basic RTF stripping - remove control words
        import re
        # Remove RTF control words
        text = re.sub(r'\\[a-z]+\d*\s?', ' ', content)
        # Remove braces
        text = re.sub(r'[{}]', '', text)
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    @staticmethod
    def load_markdown(file_path: str) -> str:
        """Load Markdown file."""
        # For markdown, we keep the formatting
        return DocumentLoader.load_text(file_path)
    
    @classmethod
    def load_file(cls, file_path: str) -> tuple[str, str]:
        """
        Load any supported file format and return (content, format).
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (text_content, file_format)
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        loaders = {
            '.txt': ('text', cls.load_text),
            '.md': ('markdown', cls.load_markdown),
            '.pdf': ('pdf', cls.load_pdf),
            '.docx': ('docx', cls.load_docx),
            '.doc': ('docx', cls.load_docx),  # Try docx loader for .doc
            '.html': ('html', cls.load_html),
            '.htm': ('html', cls.load_html),
            '.csv': ('csv', cls.load_csv),
            '.json': ('json', cls.load_json),
            '.xml': ('xml', cls.load_xml),
            '.rtf': ('rtf', cls.load_rtf),
        }
        
        if extension not in loaders:
            raise ValueError(f"Unsupported file format: {extension}")
        
        format_name, loader_func = loaders[extension]
        content = loader_func(file_path)
        
        return content, format_name
    
    @classmethod
    def load_from_bytes(cls, file_bytes: bytes, filename: str) -> tuple[str, str]:
        """
        Load file from bytes (for uploaded files).
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (text_content, file_format)
        """
        import tempfile
        
        # Save to temp file and load
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        try:
            content, format_name = cls.load_file(tmp_path)
            return content, format_name
        finally:
            # Clean up temp file
            Path(tmp_path).unlink(missing_ok=True)
    
    @staticmethod
    def get_supported_extensions() -> list[str]:
        """Get list of supported file extensions."""
        return ['.txt', '.md', '.pdf', '.docx', '.doc', '.html', '.htm', 
                '.csv', '.json', '.xml', '.rtf']
    
    @staticmethod
    def get_supported_extensions_display() -> str:
        """Get human-readable list of supported formats."""
        return "txt, md, pdf, docx, doc, html, htm, csv, json, xml, rtf"
